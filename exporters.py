"""
Exporters: PDF (via reportlab) and DOCX (via python-docx).
Each function receives an Exam ORM object and returns raw bytes.
"""
import io
from models import Exam, ExerciseType


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _render_exercise_text(exercise) -> list[str]:
    """Returns a list of text lines describing the exercise content."""
    lines = []
    c = exercise.content or {}
    etype = exercise.exercise_type

    if etype == ExerciseType.multiple_choice:
        lines.append(c.get("question", ""))
        for i, opt in enumerate(c.get("options", []), 1):
            lines.append(f"   {chr(64+i)}) {opt}")

    elif etype == ExerciseType.fill_in_the_blank:
        lines.append(c.get("sentence", ""))

    elif etype == ExerciseType.error_correction:
        lines.append(c.get("sentence", ""))
        lines.append("Correction: _______________________________")

    elif etype == ExerciseType.sentence_transformation:
        lines.append(c.get("original", ""))
        key = c.get("keyword", "")
        lines.append(f"Use the word: '{key}'")
        lines.append("Answer: _______________________________")

    elif etype == ExerciseType.true_false:
        lines.append(c.get("statement", ""))
        lines.append("True  /  False")

    else:
        # Generic fallback
        for k, v in c.items():
            lines.append(f"{k}: {v}")

    return lines


# ─────────────────────────────────────────────────────────────────────────────
# PDF
# ─────────────────────────────────────────────────────────────────────────────

def export_to_pdf(exam: Exam) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ExamTitle", parent=styles["Title"],
        fontSize=20, spaceAfter=4, textColor=colors.HexColor("#1a1a2e")
    )
    subtitle_style = ParagraphStyle(
        "Subtitle", parent=styles["Normal"],
        fontSize=11, textColor=colors.HexColor("#555"), spaceAfter=2
    )
    section_style = ParagraphStyle(
        "Section", parent=styles["Heading2"],
        fontSize=12, textColor=colors.HexColor("#2563eb"),
        spaceBefore=14, spaceAfter=4
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontSize=10, leading=15
    )
    small_style = ParagraphStyle(
        "Small", parent=styles["Normal"],
        fontSize=9, textColor=colors.HexColor("#777"), leading=13
    )

    story = []

    # Header
    story.append(Paragraph(exam.title, title_style))
    if exam.level:
        story.append(Paragraph(f"Level: {exam.level}", subtitle_style))
    if exam.description:
        story.append(Paragraph(exam.description, subtitle_style))

    # Name / Date line
    story.append(Spacer(1, 0.3*cm))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#2563eb")))
    story.append(Spacer(1, 0.3*cm))

    data = [["Name: ___________________________", "Date: _______________", "Score: _______"]]
    t = Table(data, colWidths=[9*cm, 5*cm, 4*cm])
    t.setStyle(TableStyle([
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#333")),
    ]))
    story.append(t)
    story.append(Spacer(1, 0.5*cm))

    # Exercises
    for idx, ee in enumerate(exam.exam_exercises, 1):
        ex = ee.exercise
        story.append(Paragraph(
            f"Exercise {idx} — {ex.exercise_type.replace('_', ' ').title()} "
            f"<font color='#9ca3af'>({ex.category})</font>",
            section_style
        ))
        story.append(Paragraph(ex.instructions, body_style))
        story.append(Spacer(1, 0.2*cm))
        for line in _render_exercise_text(ex):
            story.append(Paragraph(line, body_style))
        story.append(Spacer(1, 0.4*cm))

    doc.build(story)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# DOCX
# ─────────────────────────────────────────────────────────────────────────────

def export_to_docx(exam: Exam) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title_p = doc.add_heading(exam.title, level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title_p.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    if exam.level:
        p = doc.add_paragraph()
        run = p.add_run(f"Level: {exam.level}")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if exam.description:
        p = doc.add_paragraph()
        run = p.add_run(exam.description)
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # Name / date row
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    row = table.rows[0]
    row.cells[0].text = "Name: ___________________________"
    row.cells[1].text = "Date: _______________"
    row.cells[2].text = "Score: _______"
    for cell in row.cells:
        for para in cell.paragraphs:
            para.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # Exercises
    for idx, ee in enumerate(exam.exam_exercises, 1):
        ex = ee.exercise
        heading = doc.add_heading(
            f"Exercise {idx} — {ex.exercise_type.replace('_', ' ').title()}",
            level=2
        )
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)

        instr = doc.add_paragraph(ex.instructions)
        instr.runs[0].font.size = Pt(10)
        instr.runs[0].font.italic = True

        for line in _render_exercise_text(ex):
            p = doc.add_paragraph(line)
            p.paragraph_format.left_indent = Cm(0.5)
            for run in p.runs:
                run.font.size = Pt(10)

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
